"""Document parsers: PDF (pymupdf), DOCX (python-docx), plain TXT."""
from __future__ import annotations

import logging
from collections import namedtuple
from pathlib import Path

from app.core.config import settings
from app.modules.document_processing.errors import UnsupportedFileType

log = logging.getLogger(__name__)


ParsedDocument = namedtuple("ParsedDocument", "text pages raw_path")


_EXT_KIND = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
}

_MIME_KIND = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}


def detect_kind(filename: str, content_type: str | None) -> str:
    """Resolve a filename (and optional MIME) to a parser key. Raises UnsupportedFileType."""
    if filename:
        suffix = Path(filename).suffix.lower()
        if suffix in _EXT_KIND:
            return _EXT_KIND[suffix]
        allowed = settings.doc_allowed_exts_list
        if allowed and suffix not in allowed:
            raise UnsupportedFileType(f"extension '{suffix}' not in {list(allowed)}")
    if content_type:
        mt = content_type.split(";")[0].strip().lower()
        if mt in _MIME_KIND:
            return _MIME_KIND[mt]
    raise UnsupportedFileType(
        f"cannot determine file kind for filename={filename!r} content_type={content_type!r}"
    )


def parse(file_path: Path, kind: str) -> ParsedDocument:
    """Dispatch to the right parser."""
    if kind == "pdf":
        return _parse_pdf(file_path)
    if kind == "docx":
        return _parse_docx(file_path)
    if kind == "txt":
        return _parse_txt(file_path)
    # Unreachable because detect_kind validates first.
    raise UnsupportedFileType(f"unknown kind: {kind}")


def _parse_pdf(p: Path) -> ParsedDocument:
    import pymupdf  # noqa: WPS433

    doc = pymupdf.open(p)
    parts: list[str] = []
    try:
        for page in doc:
            parts.append(page.get_text("text"))
        pages = doc.page_count
    finally:
        doc.close()
    text = "\n".join(parts).strip()
    return ParsedDocument(text=text, pages=pages, raw_path=p)


def _parse_docx(p: Path) -> ParsedDocument:
    import docx  # noqa: WPS433

    document = docx.Document(str(p))
    paragraphs = [par.text for par in document.paragraphs if par.text]
    text = "\n".join(paragraphs).strip()
    # DOCX has no reliable page count without rendering; per contract we return None.
    return ParsedDocument(text=text, pages=None, raw_path=p)


def _parse_txt(p: Path) -> ParsedDocument:
    text = p.read_text(encoding="utf-8", errors="replace").strip()
    return ParsedDocument(text=text, pages=None, raw_path=p)
